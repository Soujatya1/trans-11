import streamlit as st
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
import concurrent.futures
import requests
from langdetect import detect
import time
from docx.shared import RGBColor
from docx.enum.text import WD_COLOR_INDEX
import re

def split_into_sentences(text):
    """Split text into sentences using regex patterns"""
    if not text or not text.strip():
        return []
    
    # Basic sentence splitting pattern - handles common sentence endings
    sentence_pattern = r'(?<=[.!?])\s+(?=[A-Z])'
    sentences = re.split(sentence_pattern, text.strip())
    
    # Clean up sentences and filter out empty ones
    sentences = [s.strip() for s in sentences if s.strip()]
    
    return sentences

def translate_text(text, source_language, target_language):
    if not text or not text.strip():
        return text
        
    if source_language == target_language:
        return text
        
    api_url = "https://meity-auth.ulcacontrib.org/ulca/apis/v0/model/getModelsPipeline/"
    user_id = "00fe73dcb98f43f39c1c308616856405"
    ulca_api_key = "426d392042-9028-4f13-aea7-ad172f8048f8"
    
    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {ulca_api_key}",
        "userID": user_id,
        "ulcaApiKey": ulca_api_key
    }
    
    payload = {
        "pipelineTasks": [
            {
                "taskType": "translation",
                "config": {
                    "language": {
                        "sourceLanguage": source_language,
                        "targetLanguage": target_language
                    }
                }
            }
        ],
        "pipelineRequestConfig": {
            "pipelineId": "64392f96daac500b55c543cd"
        }
    }
    
    max_retries = 3
    for attempt in range(max_retries):
        try:
            response = requests.post(api_url, json=payload, headers=headers, timeout=10)
            if response.status_code == 200:
                response_data = response.json()
                service_id = response_data["pipelineResponseConfig"][0]["config"][0]["serviceId"]
                break
            else:
                if attempt < max_retries - 1:
                    time.sleep(1)
                    continue
                else:
                    return text
        except Exception as e:
            if attempt < max_retries - 1:
                time.sleep(1)
                continue
            else:
                return text
    
    compute_payload = {
        "pipelineTasks": [
            {
                "taskType": "translation",
                "config": {
                    "language": {
                        "sourceLanguage": source_language,
                        "targetLanguage": target_language
                    },
                    "serviceId": service_id
                }
            }
        ],
        "inputData": {
            "input": [
                {
                    "source": text
                }
            ]
        }
    }
    
    callback_url = response_data["pipelineInferenceAPIEndPoint"]["callbackUrl"]
    headers2 = {
        "Content-Type": "application/json",
        response_data["pipelineInferenceAPIEndPoint"]["inferenceApiKey"]["name"]: 
        response_data["pipelineInferenceAPIEndPoint"]["inferenceApiKey"]["value"]
    }
    
    for attempt in range(max_retries):
        try:
            compute_response = requests.post(callback_url, json=compute_payload, headers=headers2, timeout=15)
            if compute_response.status_code == 200:
                compute_response_data = compute_response.json()
                translated_content = compute_response_data["pipelineResponse"][0]["output"][0]["target"]
                return translated_content
            else:
                if attempt < max_retries - 1:
                    time.sleep(1)
                    continue
                else:
                    return text
        except Exception as e:
            if attempt < max_retries - 1:
                time.sleep(1)
                continue
            else:
                return text

def detect_language(text, valid_languages):
    if not text or not text.strip():
        return "en"
    
    try:
        detected = detect(text.strip())
        if detected in valid_languages:
            return detected
        return "en"
    except Exception as e:
        print(f"Error detecting language: {e}")
        return "en"

def translate_paragraph_sentences(paragraph, target_language, valid_languages):
    """Translate a paragraph at sentence level"""
    if not paragraph.text.strip():
        return False
    
    # Split paragraph into sentences
    sentences = split_into_sentences(paragraph.text)
    if not sentences:
        return False
    
    translated_sentences = []
    translation_success = True
    
    for sentence in sentences:
        if not sentence.strip():
            translated_sentences.append(sentence)
            continue
            
        try:
            source_lang = detect_language(sentence, valid_languages)
            if source_lang == target_language:
                translated_sentences.append(sentence)
                continue
                
            translated_sentence = translate_text(sentence, source_lang, target_language)
            if translated_sentence and translated_sentence != sentence:
                translated_sentences.append(translated_sentence)
            else:
                translated_sentences.append(sentence)
                translation_success = False
                
        except Exception as e:
            print(f"Error translating sentence: {e}")
            translated_sentences.append(sentence)
            translation_success = False
    
    # Replace paragraph content with translated sentences
    if translated_sentences:
        # Clear existing runs
        for run in paragraph.runs:
            run.text = ""
        
        # Add translated content as a single run
        combined_text = " ".join(translated_sentences)
        new_run = paragraph.add_run(combined_text)
        
        # Highlight if any sentence failed translation
        if not translation_success:
            new_run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            
        return translation_success
    
    return False

def translate_run_sentences(run, target_language, valid_languages):
    """Translate a run at sentence level"""
    if not run.text.strip():
        return True
    
    sentences = split_into_sentences(run.text)
    if not sentences:
        return True
    
    translated_sentences = []
    translation_success = True
    
    for sentence in sentences:
        if not sentence.strip():
            translated_sentences.append(sentence)
            continue
            
        try:
            source_lang = detect_language(sentence, valid_languages)
            if source_lang == target_language:
                translated_sentences.append(sentence)
                continue
                
            translated_sentence = translate_text(sentence, source_lang, target_language)
            if translated_sentence and translated_sentence != sentence:
                translated_sentences.append(translated_sentence)
            else:
                translated_sentences.append(sentence)
                translation_success = False
                
        except Exception as e:
            print(f"Error translating sentence in run: {e}")
            translated_sentences.append(sentence)
            translation_success = False
    
    # Replace run content with translated sentences
    if translated_sentences:
        run.text = " ".join(translated_sentences)
        
        # Highlight if any sentence failed translation
        if not translation_success:
            run.font.highlight_color = WD_COLOR_INDEX.YELLOW
            
    return translation_success
        
def translate_doc(doc, target_language, valid_languages=None):
    if valid_languages is None:
        valid_languages = [
            "en", "ks", "ne", "bn", "mr", "sd", "te", "gu", "gom", "ur",
            "sat", "kn", "ml", "mni", "ta", "hi", "pa", "or", "doi", 
            "as", "sa", "brx", "mai"
        ]
        
    stats = {
        "sentences_processed": 0,
        "paragraphs_processed": 0,
        "successful_translations": 0,
        "failed_translations": 0
    }
    
    # Process paragraphs
    for p in doc.paragraphs:
        if p.text.strip():
            stats["paragraphs_processed"] += 1
            
            # Count sentences in this paragraph
            sentences = split_into_sentences(p.text)
            stats["sentences_processed"] += len(sentences)
            
            if len(p.runs) > 1:
                # Handle multi-run paragraphs at sentence level
                success = translate_paragraph_sentences(p, target_language, valid_languages)
                if success:
                    stats["successful_translations"] += len(sentences)
                else:
                    stats["failed_translations"] += len(sentences)
            else:
                # Handle single run paragraphs
                for run in p.runs:
                    if run.text.strip():
                        success = translate_run_sentences(run, target_language, valid_languages)
                        if success:
                            stats["successful_translations"] += len(sentences)
                        else:
                            stats["failed_translations"] += len(sentences)
   
    # Process tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    if para.text.strip():
                        stats["paragraphs_processed"] += 1
                        
                        # Count sentences in this cell paragraph
                        sentences = split_into_sentences(para.text)
                        stats["sentences_processed"] += len(sentences)
                        
                        if len(para.runs) > 1:
                            success = translate_paragraph_sentences(para, target_language, valid_languages)
                            if success:
                                stats["successful_translations"] += len(sentences)
                            else:
                                stats["failed_translations"] += len(sentences)
                        else:
                            for run in para.runs:
                                if run.text.strip():
                                    success = translate_run_sentences(run, target_language, valid_languages)
                                    if success:
                                        stats["successful_translations"] += len(sentences)
                                    else:
                                        stats["failed_translations"] += len(sentences)
                                
    return doc, stats
    
def main():
    st.title("Sentence-Level Multilingual Document Translator")
    st.write("This app translates documents at the sentence level for better translation accuracy.")
    
    uploaded_file = st.file_uploader("Upload a Word Document", type=["docx"])
    if uploaded_file:
        doc = Document(uploaded_file)
        
        language_options = {
            "English": "en",
            "Kashmiri": "ks",
            "Nepali": "ne",
            "Bengali": "bn",
            "Marathi": "mr",
            "Sindhi": "sd",
            "Telugu": "te",
            "Gujarati": "gu",
            "Gom": "gom",
            "Urdu": "ur",
            "Santali": "sat",
            "Kannada": "kn",
            "Malayalam": "ml",
            "Manipuri": "mni",
            "Tamil": "ta",
            "Hindi": "hi",
            "Punjabi": "pa",
            "Odia": "or",
            "Dogri": "doi",
            "Assamese": "as",
            "Sanskrit": "sa",
            "Bodo": "brx",
            "Maithili": "mai"
        }
        
        valid_language_codes = list(language_options.values())
        valid_language_codes.append("en")
        
        target_language = st.selectbox("Select Target Language", options=list(language_options.keys()))
        language_code = language_options[target_language]
        
        with st.expander("Advanced Options"):
            show_detected = st.checkbox("Show detected languages in document", value=True)
            show_stats = st.checkbox("Show translation statistics", value=True)
        
        if st.button("Translate Document"):
            progress_bar = st.progress(0)
            status_text = st.empty()
            
            if show_detected:
                status_text.text("Analyzing document languages...")
                progress_bar.progress(10)
                
                languages_detected = {}
                
                for p in doc.paragraphs:
                    if p.text.strip():
                        sentences = split_into_sentences(p.text)
                        for sentence in sentences:
                            if sentence.strip():
                                try:
                                    detected_lang = detect_language(sentence.strip(), valid_language_codes)
                                    if detected_lang in languages_detected:
                                        languages_detected[detected_lang] += 1
                                    else:
                                        languages_detected[detected_lang] = 1
                                except:
                                    pass
                
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                sentences = split_into_sentences(cell.text)
                                for sentence in sentences:
                                    if sentence.strip():
                                        try:
                                            detected_lang = detect_language(sentence.strip(), valid_language_codes)
                                            if detected_lang in languages_detected:
                                                languages_detected[detected_lang] += 1
                                            else:
                                                languages_detected[detected_lang] = 1
                                        except:
                                            pass
                
                progress_bar.progress(30)
                
                if languages_detected:
                    st.info("Languages detected in document (sentence level):")
                    for lang, count in sorted(languages_detected.items(), key=lambda x: x[1], reverse=True):
                        st.write(f"- {lang}: {count} sentences")
                else:
                    st.warning("Could not detect any language. Will use English as default source.")
            
            status_text.text("Translating document at sentence level... This may take several minutes for large documents.")
            progress_bar.progress(40)
            
            translated_doc, stats = translate_doc(doc, language_code, valid_language_codes)
            
            progress_bar.progress(90)
            status_text.text("Saving translated document...")
            
            with open("translated_document_sentences.docx", "wb") as f:
                translated_doc.save(f)
            
            progress_bar.progress(100)
            
            if show_stats:
                st.subheader("Translation Statistics (Sentence Level)")
                st.write(f"- Sentences processed: {stats['sentences_processed']}")
                st.write(f"- Paragraphs processed: {stats['paragraphs_processed']}")
                st.write(f"- Successful sentence translations: {stats['successful_translations']}")
                st.write(f"- Failed sentence translations: {stats['failed_translations']}")
                
                success_rate = 0
                if stats['sentences_processed'] > 0:
                    success_rate = (stats['successful_translations'] / stats['sentences_processed']) * 100
                st.write(f"- Success rate: {success_rate:.1f}%")
            
            with open("translated_document_sentences.docx", "rb") as f:
                st.download_button(
                    label="Download Translated Document",
                    data=f,
                    file_name="translated_document_sentences.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            status_text.text("")
            st.success("Sentence-level translation complete!")

if __name__ == '__main__':
    main()
